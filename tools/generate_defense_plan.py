from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).resolve().parents[1] / "Diploma_Defense_Preparation_Plan.docx"

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"
COLOR_NAVY = "0B2545"
COLOR_BLUE = "2E74B5"
COLOR_DARK_BLUE = "1F4D78"
COLOR_MUTED = "64748B"
COLOR_HEADER_FILL = "E8EEF5"
COLOR_BORDER = "CBD5E1"
COLOR_P0_FILL = "FCE8E6"
COLOR_P1_FILL = "FFF4CE"
COLOR_P2_FILL = "E8EEF5"
COLOR_OK_FILL = "E8F5E9"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM_DXA = 80
CELL_MARGIN_START_END_DXA = 120


def set_run_font(
    run,
    *,
    name=FONT_BODY,
    size=11,
    color="000000",
    bold=False,
    italic=False,
):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), name)


def configure_paragraph(paragraph, *, before=0, after=6, line_spacing=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, COLOR_BLUE, 18, 10),
        "Heading 2": (13, COLOR_BLUE, 14, 7),
        "Heading 3": (12, COLOR_DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT_BODY
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0


def set_cell_margins(cell):
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)

    values = {
        "top": CELL_MARGIN_TOP_BOTTOM_DXA,
        "bottom": CELL_MARGIN_TOP_BOTTOM_DXA,
        "start": CELL_MARGIN_START_END_DXA,
        "end": CELL_MARGIN_START_END_DXA,
    }
    for name, value in values.items():
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_table_borders(table):
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), COLOR_BORDER)


def set_repeat_table_header(row):
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {PAGE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    existing_grid = table._tbl.tblGrid
    if existing_grid is not None:
        table._tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for column_width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(column_width))
        grid.append(grid_column)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)

    set_table_borders(table)


def set_cell_text(
    cell,
    text,
    *,
    bold=False,
    color="000000",
    align=WD_ALIGN_PARAGRAPH.LEFT,
    size=9.5,
    font=FONT_BODY,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    configure_paragraph(paragraph, before=0, after=0, line_spacing=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, name=font, size=size, color=color, bold=bold)


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: COLOR_BLUE, 2: COLOR_BLUE, 3: COLOR_DARK_BLUE}[level],
        bold=True,
    )
    return paragraph


def add_body(document, text, *, bold_lead=None, color="000000"):
    paragraph = document.add_paragraph()
    configure_paragraph(paragraph)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=color)
    return paragraph


def add_metadata_row(document, label, value):
    paragraph = document.add_paragraph()
    configure_paragraph(paragraph, before=0, after=2, line_spacing=1.1)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, bold=True, color=COLOR_NAVY)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.5, color="000000")


def add_status_table(document):
    headers = ["Компонент", "Фактичний стан у репозиторії", "Оцінка", "Ризик для захисту"]
    rows = [
        (
            "Вебзастосунок",
            "React/Vite: головна сторінка, контакти та три калькулятори. Розрахунки виконуються локально у frontend, без виклику FastAPI.",
            "Частково",
            "Формули frontend і backend можуть давати різні результати; немає наскрізного сценарію створення заявки.",
        ),
        (
            "Backend API",
            "FastAPI + SQLAlchemy. Підключено clients, orders і calculator; є health-check та перевірка БД.",
            "Основа є",
            "Модулі auth, analytics і logs не підключені або не реалізовані як завершені маршрути.",
        ),
        (
            "Desktop CRM",
            "PySide6, ролі admin/manager, навігація, екрани замовлень, клієнтів і звітів.",
            "Макет",
            "ApiClient повертає порожні дані, авторизація локальна, таблиці та частина розділів не працюють з API.",
        ),
        (
            "Якість і документація",
            "Автоматичні тести не виявлені; README frontend залишився шаблонним; відтворювані інструкції запуску відсутні.",
            "Критично",
            "Збій під час демонстрації або питання комісії про перевірку якості буде складно закрити доказами.",
        ),
    ]

    table = document.add_table(rows=1, cols=4)
    set_table_geometry(table, [1600, 3460, 1000, 3300])
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=COLOR_NAVY, size=9)
        shade_cell(table.rows[0].cells[index], COLOR_HEADER_FILL)
    set_repeat_table_header(table.rows[0])

    for component, state, assessment, risk in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], component, bold=True, color=COLOR_NAVY)
        set_cell_text(cells[1], state)
        set_cell_text(cells[2], assessment, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], risk)
        if assessment == "Критично":
            shade_cell(cells[2], COLOR_P0_FILL)
        elif assessment == "Макет":
            shade_cell(cells[2], COLOR_P1_FILL)
        else:
            shade_cell(cells[2], COLOR_P2_FILL)

    document.add_paragraph()


def add_checklist_table(document, tasks):
    table = document.add_table(rows=1, cols=4)
    set_table_geometry(table, [700, 3900, 4060, 700])
    headers = ["Пріор.", "Завдання", "Критерій готовності", "Статус"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=COLOR_NAVY, size=9)
        shade_cell(table.rows[0].cells[index], COLOR_HEADER_FILL)
    set_repeat_table_header(table.rows[0])

    for priority, task, criterion in tasks:
        cells = table.add_row().cells
        set_cell_text(
            cells[0],
            priority,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=COLOR_NAVY,
        )
        set_cell_text(cells[1], task)
        set_cell_text(cells[2], criterion)
        set_cell_text(cells[3], "[ ]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        shade_cell(
            cells[0],
            {"P0": COLOR_P0_FILL, "P1": COLOR_P1_FILL, "P2": COLOR_P2_FILL}[priority],
        )

    document.add_paragraph()


def add_sequence_table(document):
    rows = [
        (
            "1. Стабілізувати",
            "Зафіксувати команди запуску, залежності, конфігурацію БД та один набір демонстраційних даних.",
            "Усі три частини стартують на чистому середовищі за письмовою інструкцією.",
        ),
        (
            "2. З'єднати",
            "Прибрати розриви між frontend, FastAPI, БД і desktop; обрати одну точку істини для розрахунків.",
            "Один наскрізний сценарій проходить без ручного перенесення даних.",
        ),
        (
            "3. Довести",
            "Додати мінімальний набір автоматичних тестів і чек-лист ручної перевірки.",
            "Є відтворювані докази коректності ключових формул та CRUD-операцій.",
        ),
        (
            "4. Пояснити",
            "Підготувати схеми архітектури й БД, слайди та короткий сценарій демонстрації.",
            "Виступ вкладається у ліміт часу, а кожне технічне рішення має коротке обґрунтування.",
        ),
    ]
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [1600, 4260, 3500])
    headers = ["Етап", "Що робимо", "Результат"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=COLOR_NAVY)
        shade_cell(table.rows[0].cells[index], COLOR_HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    for stage, action, result in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], stage, bold=True, color=COLOR_NAVY)
        set_cell_text(cells[1], action)
        set_cell_text(cells[2], result)
    document.add_paragraph()


def add_risk_table(document):
    rows = [
        (
            "Різні формули ціни",
            "Frontend і backend мають окремі константи та алгоритми.",
            "Залишити розрахунок у backend; frontend лише надсилає параметри та показує відповідь.",
            "P0",
        ),
        (
            "Desktop показує порожні дані",
            "ApiClient містить заглушки, таблиці не завантажують API.",
            "Реалізувати read-only отримання нових замовлень, клієнтів і статистики; потім CRUD.",
            "P0",
        ),
        (
            "Демонстрація залежить від середовища",
            "Немає кореневого README, є кілька venv, requirements для desktop видалено у робочій копії.",
            "Створити єдину інструкцію запуску, зафіксувати залежності, додати .env.example без секретів.",
            "P0",
        ),
        (
            "Незавершені розділи CRM",
            "Analytics/users/logs ведуть на PlaceholderPage; reports містить порожню таблицю.",
            "Не розширювати обсяг без потреби: завершити один переконливий звіт і приховати недемонстраційні пункти.",
            "P1",
        ),
        (
            "Немає доказів якості",
            "Тести не знайдені.",
            "Додати unit-тести калькуляторів, API smoke-тести і сценарій ручної перевірки UI.",
            "P1",
        ),
    ]
    table = document.add_table(rows=1, cols=4)
    set_table_geometry(table, [1700, 2860, 3900, 900])
    headers = ["Ризик", "Ознака", "Міра", "Пріор."]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=COLOR_NAVY)
        shade_cell(table.rows[0].cells[index], COLOR_HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    for risk, evidence, mitigation, priority in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], risk, bold=True, color=COLOR_NAVY)
        set_cell_text(cells[1], evidence)
        set_cell_text(cells[2], mitigation)
        set_cell_text(cells[3], priority, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cells[3], COLOR_P0_FILL if priority == "P0" else COLOR_P1_FILL)
    document.add_paragraph()


def add_demo_table(document):
    rows = [
        (
            "1",
            "Відвідувач відкриває сайт і обирає рекламну послугу.",
            "React-маршрути та зрозумілий каталог послуг.",
        ),
        (
            "2",
            "Вводить параметри білборда, LED або друку й отримує розрахунок.",
            "Єдина формула у FastAPI, валідація вхідних даних, відповідь API.",
        ),
        (
            "3",
            "Надсилає контактні дані або заявку.",
            "Заявка зберігається у БД і має ідентифікатор/статус.",
        ),
        (
            "4",
            "Менеджер входить у desktop CRM і бачить нову заявку.",
            "Рольовий доступ, ApiClient, завантаження даних із backend.",
        ),
        (
            "5",
            "Менеджер створює/оновлює клієнта та замовлення.",
            "CRUD clients/orders, перерахунок ПДВ/знижки, зміна статусу.",
        ),
        (
            "6",
            "Формує короткий звіт або показує підсумкові показники.",
            "Один завершений звіт або аналітична картка на реальних даних.",
        ),
    ]
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [700, 4430, 4230])
    headers = ["Крок", "Дія під час показу", "Що доводить"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=COLOR_NAVY)
        shade_cell(table.rows[0].cells[index], COLOR_HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    for step, action, evidence in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], step, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], action)
        set_cell_text(cells[2], evidence)
    document.add_paragraph()


def add_definition_of_done(document):
    rows = [
        ("Запуск", "Backend, frontend і desktop стартують за README на іншому ПК або в чистому середовищі."),
        ("Дані", "Є демонстраційний набір: користувач, клієнт, мінімум три замовлення різних типів."),
        ("Інтеграція", "Калькулятор, заявки, клієнти та замовлення використовують backend/БД без ручного дублювання."),
        ("Надійність", "Критичні unit/API тести проходять; ручний smoke-чек зафіксовано."),
        ("Безпека", "Секрети не закомічені; паролі не зберігаються у відкритому вигляді у фінальній версії."),
        ("Документація", "Є схема архітектури, ER-діаграма, опис ролей, алгоритмів розрахунку та інструкція запуску."),
        ("Захист", "Слайди готові, демонстрація до 5-7 хвилин, є резервні скріншоти/відео та відповіді на типові питання."),
    ]
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 6760, 700])
    headers = ["Область", "Умова завершення", "Готово"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=COLOR_NAVY)
        shade_cell(table.rows[0].cells[index], COLOR_HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    for area, condition in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], area, bold=True, color=COLOR_NAVY)
        set_cell_text(cells[1], condition)
        set_cell_text(cells[2], "[ ]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    document.add_paragraph()


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    configure_paragraph(paragraph, before=0, after=0, line_spacing=1.0)
    prefix = paragraph.add_run("Сторінка ")
    set_run_font(prefix, size=9, color=COLOR_MUTED)
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(page_field)


def add_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    configure_paragraph(paragraph, before=0, after=0, line_spacing=1.0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = paragraph.add_run("Підготовка до захисту диплома")
    set_run_font(left, size=9, color=COLOR_MUTED, bold=True)
    right = paragraph.add_run("\tCreative Spark Agency")
    set_run_font(right, size=9, color=COLOR_MUTED)


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(document)
    add_header(section)
    add_footer(section)

    kicker = document.add_paragraph()
    configure_paragraph(kicker, before=10, after=4, line_spacing=1.0)
    kicker_run = kicker.add_run("ТЕХНІЧНИЙ ПЛАН ПІДГОТОВКИ")
    set_run_font(kicker_run, size=10, color=COLOR_BLUE, bold=True)

    title = document.add_paragraph()
    configure_paragraph(title, before=0, after=5, line_spacing=1.0)
    title_run = title.add_run("Підготовка до захисту дипломного проєкту")
    set_run_font(title_run, size=25, color=COLOR_NAVY, bold=True)

    subtitle = document.add_paragraph()
    configure_paragraph(subtitle, before=0, after=16, line_spacing=1.15)
    subtitle_run = subtitle.add_run(
        "Інформаційна система рекламного агентства: вебсайт, FastAPI backend і desktop CRM"
    )
    set_run_font(subtitle_run, size=13, color=COLOR_MUTED)

    add_metadata_row(document, "Горизонт", "1-2 тижні")
    add_metadata_row(document, "Головний фокус", "технічна якість і надійність демонстрації")
    add_metadata_row(document, "Формат", "чек-лист за робочими блоками")
    add_metadata_row(document, "Джерело", "локальний репозиторій advertising_agency_project")
    add_metadata_row(document, "Актуально на", "7 червня 2026 року")

    document.add_paragraph()
    lead = document.add_paragraph()
    configure_paragraph(lead, before=4, after=10, line_spacing=1.2)
    lead_run = lead.add_run(
        "Мета: за короткий термін перетворити наявні окремі компоненти на відтворюваний "
        "наскрізний сценарій, який стабільно запускається, працює з єдиними даними та має докази якості."
    )
    set_run_font(lead_run, size=11.5, color=COLOR_NAVY, bold=True)

    add_heading(document, "1. Поточна готовність", level=1)
    add_body(
        document,
        "Оцінка нижче базується на файлах і маршрутах, які є в репозиторії. "
        "Вона потрібна для визначення обсягу робіт, а не як остаточний аудит усієї системи.",
    )
    add_status_table(document)

    add_heading(document, "2. Рекомендований порядок", level=1)
    add_body(
        document,
        "Не починайте з презентації або нових функцій. Спочатку забезпечте стабільний запуск і "
        "один завершений бізнес-сценарій, потім додайте тести й матеріали захисту.",
    )
    add_sequence_table(document)

    document.add_page_break()
    add_heading(document, "3. Чек-лист технічної підготовки", level=1)

    add_heading(document, "3.1. Відтворюваний запуск і конфігурація", level=2)
    add_checklist_table(
        document,
        [
            (
                "P0",
                "Створити кореневий README з окремими командами запуску backend, frontend і desktop.",
                "Нова людина запускає систему без усних підказок; описані порти, порядок запуску та тестові облікові дані.",
            ),
            (
                "P0",
                "Зафіксувати Python- і Node-залежності та обрати одну актуальну версію Python/venv.",
                "Є requirements/pyproject для Python-компонентів і package-lock для frontend; зайві середовища не потрібні для інструкції.",
            ),
            (
                "P0",
                "Додати .env.example без паролів і перевірити, що реальний .env не потрапляє у Git.",
                "Усі потрібні змінні перелічені; секрети відсутні у репозиторії та демонстраційних матеріалах.",
            ),
            (
                "P1",
                "Налаштувати .gitignore для __pycache__, venv, build/dist і локальних файлів БД/логів.",
                "Після запуску застосунків git status не заповнюється згенерованими файлами.",
            ),
            (
                "P1",
                "Підготувати seed-скрипт або дамп демонстраційних даних.",
                "Однією командою створюються користувач, клієнти, послуги й замовлення для показу.",
            ),
        ],
    )

    add_heading(document, "3.2. Єдина архітектура та інтеграція", level=2)
    add_checklist_table(
        document,
        [
            (
                "P0",
                "Зробити FastAPI єдиною точкою істини для розрахунку вартості.",
                "Frontend надсилає параметри до /calculator/*; локальні дублікати формул видалені або використовуються лише як явно позначений fallback.",
            ),
            (
                "P0",
                "Узгодити довідники й назви полів для billboard, LED і print між frontend-схемами та backend-схемами.",
                "Однаковий набір параметрів дає один результат; помилки валідації показуються користувачу.",
            ),
            (
                "P0",
                "Реалізувати реальні HTTP-виклики в desktop/services/api_client.py.",
                "CRM отримує клієнтів, замовлення та статистику з FastAPI, обробляє timeout/помилки й показує зрозуміле повідомлення.",
            ),
            (
                "P1",
                "Підключити CORS і централізовану конфігурацію base URL.",
                "Frontend та desktop працюють з одним API URL без зміни коду перед показом.",
            ),
            (
                "P1",
                "Перевірити реєстрацію всіх потрібних router у backend/main.py.",
                "У Swagger видно лише підтримувані маршрути; порожні або незавершені auth/analytics/logs не створюють хибних очікувань.",
            ),
        ],
    )

    add_heading(document, "3.3. Наскрізний бізнес-сценарій", level=2)
    add_checklist_table(
        document,
        [
            (
                "P0",
                "Перетворити форму контакту/замовлення на збереження заявки у backend.",
                "Після відправлення запис з'являється у БД і доступний у desktop CRM як нова заявка.",
            ),
            (
                "P0",
                "Завершити відображення списку клієнтів і замовлень у desktop.",
                "Таблиці завантажують реальні дані, підтримують оновлення і коректно показують порожній/помилковий стан.",
            ),
            (
                "P0",
                "Реалізувати мінімальні create/update/status операції для замовлення.",
                "Менеджер створює або редагує замовлення, змінює статус, а суми ПДВ/знижки перераховуються backend.",
            ),
            (
                "P1",
                "Замінити локальний словник паролів на backend-авторизацію або чітко обмежити демонстраційний режим.",
                "Паролі не зберігаються у відкритому вигляді; роль admin/manager перевіряється в одному місці.",
            ),
            (
                "P1",
                "Завершити один демонстраційний звіт замість кількох порожніх розділів.",
                "Звіт фільтрується за періодом/типом і показує суму продажів, ПДВ та знижки на реальних даних.",
            ),
            (
                "P2",
                "Приховати або позначити як майбутню роботу розділи analytics/users/logs, якщо вони не входять у захист.",
                "Комісія не переходить у порожні екрани під час основного сценарію.",
            ),
        ],
    )

    add_heading(document, "3.4. Тести та перевірка якості", level=2)
    add_checklist_table(
        document,
        [
            (
                "P0",
                "Додати unit-тести формул billboard, LED, print, ПДВ і знижки.",
                "Перевірені звичайні, граничні та некоректні значення; очікувані суми зафіксовані.",
            ),
            (
                "P1",
                "Додати API smoke-тести для /health, /calculator, /clients і /orders.",
                "Тести створюють дані, читають їх, змінюють статус і перевіряють коди 200/201/204/404.",
            ),
            (
                "P1",
                "Запускати frontend lint і production build.",
                "npm run lint і npm run build завершуються без помилок перед кожною репетицією.",
            ),
            (
                "P1",
                "Додати import/smoke-перевірку desktop без ручного кліку по всіх екранах.",
                "Головні модулі імпортуються, ApiClient проходить контрольний запит, UI відкривається без traceback.",
            ),
            (
                "P1",
                "Скласти ручний регресійний чек-лист з 10-15 коротких перевірок.",
                "Перед показом за 10 хвилин можна перевірити запуск, калькуляцію, заявку, CRUD і звіт.",
            ),
        ],
    )

    add_heading(document, "3.5. Дані, безпека і відмовостійкість", level=2)
    add_checklist_table(
        document,
        [
            (
                "P0",
                "Перевірити підключення до БД і поведінку при недоступній БД.",
                "/health коректно показує стан; UI не зависає і пояснює помилку користувачу.",
            ),
            (
                "P1",
                "Додати валідацію дат, кількості, від'ємних сум і невідомих довідникових значень.",
                "Некоректні запити повертають контрольовані 4xx, а UI не показує NaN або неправильну ціну.",
            ),
            (
                "P1",
                "Перевірити цілісність видалення клієнтів і пов'язаних замовлень.",
                "Система не втрачає пов'язані дані; помилка пояснює, чому видалення неможливе.",
            ),
            (
                "P1",
                "Підготувати резервний сценарій демонстрації.",
                "Є локальна БД/seed, скріншоти ключових екранів і коротке відео на випадок проблем із середовищем.",
            ),
        ],
    )

    document.add_page_break()
    add_heading(document, "4. Матеріали для пояснювальної записки", level=1)
    add_checklist_table(
        document,
        [
            (
                "P0",
                "Намалювати контекстну й компонентну схеми.",
                "Показані користувач, browser frontend, FastAPI, desktop CRM, SQL database і напрямки обміну.",
            ),
            (
                "P0",
                "Підготувати ER-діаграму.",
                "Відображені User, Client, Order, Service, Payment, Expense, AdvertisingSpace і ключові зв'язки.",
            ),
            (
                "P1",
                "Описати три алгоритми калькуляції та фінансовий розрахунок замовлення.",
                "Для кожного алгоритму наведені вхідні параметри, коефіцієнти, формула, округлення і приклад.",
            ),
            (
                "P1",
                "Описати розподіл відповідальності між компонентами.",
                "Frontend відповідає за взаємодію, backend за правила/валідацію, БД за зберігання, desktop за робоче місце менеджера.",
            ),
            (
                "P1",
                "Зафіксувати обмеження та майбутній розвиток.",
                "Чесно перелічені незавершені модулі, production-безпека, міграції, контейнеризація та розширення аналітики.",
            ),
            (
                "P2",
                "Зібрати таблицю використаних технологій з обґрунтуванням.",
                "Для React, FastAPI, SQLAlchemy, PySide6 і СУБД є 1-2 аргументи, пов'язані з вимогами проєкту.",
            ),
        ],
    )

    add_heading(document, "5. Реєстр основних ризиків", level=1)
    add_risk_table(document)

    add_heading(document, "6. Сценарій демонстрації", level=1)
    add_body(
        document,
        "Оптимальний показ має бути наскрізним і тривати 5-7 хвилин. Не демонструйте кожну сторінку; "
        "показуйте, як одна заявка проходить через систему.",
    )
    add_demo_table(document)

    add_heading(document, "7. Питання, до яких варто підготувати відповіді", level=1)
    questions = [
        "Чому обрано окремі web і desktop клієнти, а не один інтерфейс?",
        "Де зберігається бізнес-логіка і як уникнуто різних результатів калькуляції?",
        "Як організовані зв'язки між клієнтом, замовленням, послугою, оплатою та витратами?",
        "Як перевіряється коректність ПДВ, знижок, дат і кількості показів/тиражу?",
        "Як система поводиться при недоступній БД або API?",
        "Які права мають admin і manager та як ці права контролюються?",
        "Які тести виконані і які дефекти вони можуть виявити?",
        "Що потрібно змінити перед production-розгортанням?",
    ]
    question_table = document.add_table(rows=1, cols=2)
    set_table_geometry(question_table, [700, 8660])
    set_cell_text(question_table.rows[0].cells[0], "№", bold=True, color=COLOR_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(question_table.rows[0].cells[1], "Питання", bold=True, color=COLOR_NAVY)
    shade_cell(question_table.rows[0].cells[0], COLOR_HEADER_FILL)
    shade_cell(question_table.rows[0].cells[1], COLOR_HEADER_FILL)
    set_repeat_table_header(question_table.rows[0])
    for index, question in enumerate(questions, start=1):
        cells = question_table.add_row().cells
        set_cell_text(cells[0], str(index), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], question)
    document.add_paragraph()

    add_heading(document, "8. Definition of Done перед захистом", level=1)
    add_definition_of_done(document)

    add_heading(document, "9. Перша робоча сесія", level=1)
    add_body(
        document,
        "Почніть із P0-завдань у блоках 3.1-3.3. Практична ціль першої сесії: "
        "запустити всі компоненти, виконати один розрахунок через FastAPI, зберегти заявку у БД "
        "та побачити її у desktop CRM. Це найцінніший доказ цілісності дипломного проєкту.",
        bold_lead="Почніть із P0-завдань у блоках 3.1-3.3.",
    )

    final_note = document.add_paragraph()
    configure_paragraph(final_note, before=8, after=0, line_spacing=1.15)
    final_run = final_note.add_run(
        "Примітка: план складено за поточним станом локального репозиторію. "
        "Перед виконанням завдань варто узгодити їх із вимогами кафедри та фактичним регламентом захисту."
    )
    set_run_font(final_run, size=9.5, color=COLOR_MUTED, italic=True)

    document.core_properties.title = "Підготовка до захисту дипломного проєкту"
    document.core_properties.subject = "Технічний чек-лист для advertising agency project"
    document.core_properties.author = "Codex"
    document.core_properties.keywords = "диплом, захист, FastAPI, React, PySide6, чек-лист"

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
